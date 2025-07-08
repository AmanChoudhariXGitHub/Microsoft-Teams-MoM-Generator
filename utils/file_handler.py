import os
from fpdf import FPDF
from docx import Document
from config import UPLOAD_FOLDER, DOWNLOAD_FOLDER

def save_uploaded_file(uploaded_file):
    """Save uploaded file to UPLOAD_FOLDER and return file path."""
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

    file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    return file_path


def export_to_pdf(mom_text, filename="Meeting_Minutes.pdf"):
    """Export MoM text to a structured PDF file."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", style='B', size=14)
    pdf.cell(200, 10, txt="Meeting Minutes", ln=True, align="C")
    
    pdf.set_font("Arial", size=12)
    pdf.ln(10)  # Line break

    for line in mom_text.strip().split("\n"):
        pdf.multi_cell(0, 10, txt=line.strip())

    os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
    file_path = os.path.join(DOWNLOAD_FOLDER, filename)
    pdf.output(file_path)
    return file_path


def export_to_docx(mom_text, filename="Meeting_Minutes.docx"):
    """Export MoM text to a structured DOCX file."""
    doc = Document()
    doc.add_heading("Meeting Minutes", level=1)

    for line in mom_text.strip().split("\n"):
        doc.add_paragraph(line.strip())

    os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
    file_path = os.path.join(DOWNLOAD_FOLDER, filename)
    doc.save(file_path)
    return file_path
